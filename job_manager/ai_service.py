import base64
import os
from io import BytesIO
from django.conf import settings
from django.core.files.base import ContentFile
from openai import OpenAI
import markdown
from xhtml2pdf import pisa
from docx import Document

from main.models import Profile
from resume.models import Experience, Education, Skill
from project.models import Project

def get_portfolio_context():
    """Compiles all relevant portfolio data into a single context string."""
    try:
        profile = Profile.objects.first()
        profile_info = f"Name: {profile.name}\nTitle: {profile.title}\nEmail: {profile.email}\nAbout: {profile.about_text}\n" if profile else "No profile info.\n"
    except:
        profile_info = "Profile info not available.\n"
        
    skills = "\n".join([f"- {s.name} ({s.proficiency}%)" for s in Skill.objects.all()])
    experience = "\n".join([f"- {e.role} at {e.company} ({e.start_date} to {e.end_date}): {e.description}" for e in Experience.objects.all()])
    education = "\n".join([f"- {e.degree} at {e.institution} ({e.start_date} to {e.end_date}): {e.description}" for e in Education.objects.all()])
    projects = "\n".join([f"- {p.title}: {p.description}" for p in Project.objects.all()[:5]]) # Top 5 projects
    
    return f"""
    --- MY BACKGROUND ---
    {profile_info}
    SKILLS:
    {skills}
    
    EXPERIENCE:
    {experience}
    
    EDUCATION:
    {education}
    
    TOP PROJECTS:
    {projects}
    """

def generate_application_documents(job_description_text, job_image_path=None):
    """
    Calls OpenAI to generate a tailored CV and Cover Letter.
    Returns a tuple: (generated_cv, generated_cover_letter)
    """
    client = OpenAI(api_key=settings.OPENAI_API_KEY)
    
    # Base prompt
    system_prompt = "You are an expert career coach and professional resume writer."
    context_str = get_portfolio_context()
    
    content_messages = [
        {"type": "text", "text": f"Here is my background information:\n{context_str}\n\n"}
    ]
    
    if job_description_text:
        content_messages.append({"type": "text", "text": f"Here is the job description I am applying for:\n{job_description_text}\n\n"})
        
    if job_image_path and os.path.exists(job_image_path):
        with open(job_image_path, "rb") as image_file:
            base64_image = base64.b64encode(image_file.read()).decode('utf-8')
            content_messages.append({
                "type": "image_url",
                "image_url": {
                    "url": f"data:image/jpeg;base64,{base64_image}"
                }
            })
            content_messages.append({"type": "text", "text": "Extract the job requirements and role details from this screenshot of a job posting."})

    content_messages.append({
        "type": "text", 
        "text": "Based strictly on my background and the provided job description, please write:\n1. A highly tailored professional Cover Letter.\n2. A tailored CV (Resume) highlighting the most relevant skills and experiences for this specific role.\nFormat your response EXACTLY as follows:\n[COVER LETTER START]\n(cover letter text here)\n[COVER LETTER END]\n\n[CV START]\n(cv text here)\n[CV END]"
    })

    try:
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": content_messages}
            ],
            max_tokens=2500
        )
        
        result_text = response.choices[0].message.content
        
        # Parse the response
        cover_letter = ""
        cv = ""
        
        if "[COVER LETTER START]" in result_text and "[COVER LETTER END]" in result_text:
            cover_letter = result_text.split("[COVER LETTER START]")[1].split("[COVER LETTER END]")[0].strip()
            
        if "[CV START]" in result_text and "[CV END]" in result_text:
            cv = result_text.split("[CV START]")[1].split("[CV END]")[0].strip()
            
        # Fallback if AI didn't follow formatting strictly
        if not cover_letter or not cv:
            cover_letter = "Warning: Failed to parse format correctly. Raw output:\n\n" + result_text
            cv = result_text
            
        return cv, cover_letter
        
    except Exception as e:
        error_msg = f"Error generating documents via OpenAI: {str(e)}"
        return error_msg, error_msg

def create_pdf(text, filename_prefix="document"):
    html_content = markdown.markdown(text)
    styled_html = f"<html><head><style>body {{ font-family: Helvetica, Arial, sans-serif; font-size: 11pt; line-height: 1.4; }}</style></head><body>{html_content}</body></html>"
    result = BytesIO()
    pdf = pisa.pisaDocument(BytesIO(styled_html.encode("UTF-8")), result)
    if not pdf.err:
        return ContentFile(result.getvalue(), name=f"{filename_prefix}.pdf")
    return None

def create_docx(text, filename_prefix="document"):
    document = Document()
    # Simple markdown to docx converter
    for line in text.split('\n'):
        if line.startswith('# '):
            document.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            document.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            document.add_heading(line[4:], level=3)
        elif line.startswith('- '):
            document.add_paragraph(line[2:], style='List Bullet')
        elif line.startswith('**') and line.endswith('**'):
            p = document.add_paragraph()
            p.add_run(line[2:-2]).bold = True
        elif line.strip():
            document.add_paragraph(line)
            
    result = BytesIO()
    document.save(result)
    return ContentFile(result.getvalue(), name=f"{filename_prefix}.docx")

def process_job_application_task(application_id):
    from .models import JobApplication
    try:
        app = JobApplication.objects.get(id=application_id)
        if app.status != 'draft':
            return
            
        # Optional: wait briefly to ensure file is completely saved to disk
        import time
        time.sleep(1)
        
        image_path = app.job_description_image.path if app.job_description_image else None
        cv, cl = generate_application_documents(app.job_description_text, image_path)
        
        app.generated_cv = cv
        app.generated_cover_letter = cl
        
        # Generate files
        if cv:
            pdf_file = create_pdf(cv, "cv")
            if pdf_file: app.cv_pdf.save(f"CV_{app.company_name.replace(' ', '_')}.pdf", pdf_file, save=False)
            docx_file = create_docx(cv, "cv")
            if docx_file: app.cv_word.save(f"CV_{app.company_name.replace(' ', '_')}.docx", docx_file, save=False)
            
        if cl:
            cl_pdf = create_pdf(cl, "cover_letter")
            if cl_pdf: app.cover_letter_pdf.save(f"Cover_Letter_{app.company_name.replace(' ', '_')}.pdf", cl_pdf, save=False)
            cl_docx = create_docx(cl, "cover_letter")
            if cl_docx: app.cover_letter_word.save(f"Cover_Letter_{app.company_name.replace(' ', '_')}.docx", cl_docx, save=False)
            
        if "Error generating documents" not in cl:
            app.status = 'generated'
        
        app.save()
    except Exception as e:
        print(f"Error in AI background task: {e}")
